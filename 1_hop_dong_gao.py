import os
import sys
import subprocess
from datetime import datetime, timedelta

# =====================================================================
# KHỐI LỆNH TỰ ĐỘNG TRANG BỊ VŨ KHÍ 
# =====================================================================
def auto_install_vukhi(package_name):
    try:
        __import__(package_name)
    except ImportError:
        print(f"🛠 Tớ phát hiện máy cậu đang thiếu cây cọ vẽ bảng '{package_name}'...")
        print(f"⚙️ Đang tự động cài đặt luôn cho cậu. Cậu ngồi nhâm nhi cafe đợi 10 giây nhé!")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])
        print(f"✅ Trang bị vũ khí {package_name} thành công! Máy móc bắt đầu chạy...\n")

auto_install_vukhi('docxcompose')

# Sau khi đảm bảo đã có đủ vũ khí, mới gọi các thư viện ra làm việc
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =====================================================================
# HƯỚNG DẪN CẬU SỬA FILE WORD MẪU (BẢN VẼ 2 BẢNG)
# 
# 1. Phần Điều 1 (BẢNG TRÊN): 
#    Cậu XÓA TOÀN BỘ các dòng gạch đầu dòng (Tên hàng, Số lượng...) đi.
#    Ngay tại vị trí đó, cậu gõ 1 dòng text này:
#    {{ bang_dieu_1 }}
#
# 2. Phần Biên bản (BẢNG DƯỚI): 
#    Cậu XÓA TOÀN BỘ cái bảng cũ trong file Word đi.
#    Ngay tại vị trí đó, cậu gõ 1 dòng text này:
#    {{ bang_hang_hoa }}
#
# 3. Chỗ thời gian giao hàng:
#    Từ ngày {{ ngay }}/{{ thang }} đến hết ngày {{ giaohang }}/{{ thanggiaohang }}/{{ namgiaohang }}
# =====================================================================

EXCEL_FILE = 'thong_tin_hop_dong_gao.xlsx'  
WORD_TEMPLATE = 'mau_hop_dong_gao.docx' 
OUTPUT_DIR = ''

def format_number(val):
    """Trang điểm cho con số: 1000000 -> 1.000.000"""
    try:
        if pd.isna(val) or val == "":
            return ""
        return f"{float(val):,.0f}".replace(',', '.')
    except ValueError:
        return str(val)

def xu_ly_mst(val):
    """Radar quét MST thông minh cho cả Công ty, Chi nhánh, Hộ KD"""
    if pd.isna(val) or str(val).strip() == "":
        return ""
    
    s = str(val).strip()
    if s.endswith('.0'):
        s = s[:-2]
        
    if s.isdigit():
        if len(s) <= 10:
            return s.zfill(10)
        elif len(s) <= 13:
            s_13 = s.zfill(13)
            return f"{s_13[:10]}-{s_13[10:]}"
            
    return s

def format_cell(cell, text, bold=False, align='center'):
    """Nghệ nhân định dạng: Đổ chữ vào ô, ép Font Times New Roman 12, căn lề chuẩn chỉ + Spacing Before 6pt"""
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Căn giữa theo chiều dọc
    
    p = cell.paragraphs[0]
    
    # Ép Spacing Before 6pt cho tất cả các đoạn văn trong bảng theo lệnh của cậu
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    
    # Căn lề theo chiều ngang
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    # Ép chuẩn Font và Size
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold

def ve_vien_ngang(table):
    """Bùa chú OXML: Đưa bảng về No Borders -> Chỉ kẻ Inside Horizontal (loại bỏ viền trên/dưới cùng)"""
    tblPr = table._tbl.tblPr
    
    # 1. Quét sạch toàn bộ viền cũ (Thao tác: No borders)
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblBorders.clear()
    else:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
        
    # 2. Thao tác: Kẻ viền ngang bên trong (Chỉ dùng insideH đúng như ý cậu)
    for border_name in ['insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # Độ dày nét đứt chuẩn của Word
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)

def autofit_to_window(table):
    """Bùa chú OXML: Ép bảng tự động giãn căng tràn lề (AutoFit to Window 100%)"""
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    
    # 5000 theo chuẩn pct (percentage) của Word có nghĩa là 100% chiều rộng trang
    tblW.set(qn('w:w'), '5000') 
    tblW.set(qn('w:type'), 'pct')

def tao_hop_dong_gom_nhom():
    try:
        df = pd.read_excel(EXCEL_FILE)
        df = df.dropna(subset=['sohd']) 
    except Exception as e:
        print(f"❌ Úi, lỗi đọc Excel rồi cậu ơi: {e}")
        return
    
    # Lấy số hóa đơn đầu tiên và cuối cùng
    danh_sach_sohd = sorted(df['sohd'].dropna().astype(int))

    sohd_dau = str(danh_sach_sohd[0]).zfill(2)
    sohd_cuoi = str(danh_sach_sohd[-1]).zfill(2)

    # Tạo tên folder động
    output_dir = f"Hop_Dong_Gao_HD_{sohd_dau}_Den_{sohd_cuoi}"

    # Tạo folder nếu chưa tồn tại
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    print("🚀 Bắt đầu tạo hàng loạt hợp đồng...\n")

    nhom_hop_dong = df.groupby('sohd')
    
    # BIẾN MỚI: Danh sách để lưu lại đường dẫn của các file Word con vừa sinh ra
    danh_sach_file_da_tao = []

    for sohd, group in nhom_hop_dong:
        doc = DocxTemplate(WORD_TEMPLATE)
        
        first_row = group.iloc[0]
        so_hd_str = str(sohd).replace('.0', '').zfill(2)
        ngay = str(first_row['ngay']).replace('.0', '').zfill(2)
        thang = str(first_row['thang']).replace('.0', '').zfill(2)
        nam = str(first_row['nam']).replace('.0', '')
        
        # --- CỖ MÁY TÍNH TOÁN KẾ TOÁN ---
        # Tính TỔNG tiền thuế và TỔNG tiền thanh toán cho toàn bộ mặt hàng trong hợp đồng này
        tong_thue_hop_dong = group['thue'].sum(skipna=True)
        tong_tien_hop_dong = group['tongtien'].sum(skipna=True)
        
        # --- CỖ MÁY THỜI GIAN ---
        try:
            ngay_ky_obj = datetime(int(nam), int(thang), int(ngay))
            ngay_giao_obj = ngay_ky_obj + timedelta(days=2)
            giaohang = ngay_giao_obj.strftime("%d")
            thanggiaohang = ngay_giao_obj.strftime("%m")
            namgiaohang = ngay_giao_obj.strftime("%Y")
        except Exception:
            giaohang, thanggiaohang, namgiaohang = "...", "...", "..."
            
        so_mat_hang = len(group)

        # =================================================================
        # BẢNG 1: BẢNG Ở ĐIỀU 1 (CHỈ KẺ VIỀN NGANG BÊN TRONG - NHƯ ẢNH MẪU)
        # =================================================================
        bang_1_subdoc = doc.new_subdoc()
        bang_1 = bang_1_subdoc.add_table(rows=so_mat_hang + 1, cols=4)
        ve_vien_ngang(bang_1) # Kích hoạt bùa chú viền ngang Inside Horizontal
        autofit_to_window(bang_1) # Ép căng tràn trang giấy
        
        # Tiêu đề bảng 1
        hdr1 = bang_1.rows[0].cells
        format_cell(hdr1[0], 'Tên hàng', bold=True, align='center')
        format_cell(hdr1[1], 'Số lượng', bold=True, align='center')
        format_cell(hdr1[2], 'Đơn giá', bold=True, align='center')
        format_cell(hdr1[3], 'Mặt hàng', bold=True, align='center')
        
        # Điền dữ liệu bảng 1
        for i, (_, row) in enumerate(group.iterrows()):
            r1 = bang_1.rows[i + 1].cells
            format_cell(r1[0], str(row['tenhang']) if pd.notna(row['tenhang']) else "", align='center')
            format_cell(r1[1], f"{format_number(row['soluong'])} kg", align='center')
            format_cell(r1[2], f"{format_number(row['dongia'])} đ/kg", align='right') # Đơn giá căn phải
            format_cell(r1[3], str(row['mathang']) if pd.notna(row['mathang']) else "", align='center')

        # =================================================================
        # BẢNG 2: BẢNG GIAO NHẬN HÀNG HÓA (CÓ LƯỚI FULL)
        # =================================================================
        bang_subdoc = doc.new_subdoc()
        bang = bang_subdoc.add_table(rows=so_mat_hang + 3, cols=5)
        bang.style = 'Table Grid'
        autofit_to_window(bang) # Ép căng tràn trang giấy
        
        # Ghi Tiêu đề bảng 2
        hdr_cells = bang.rows[0].cells
        format_cell(hdr_cells[0], 'STT', bold=True, align='center')
        format_cell(hdr_cells[1], 'Tên hàng hóa', bold=True, align='center')
        format_cell(hdr_cells[2], 'Số lượng (Kg)', bold=True, align='center')
        format_cell(hdr_cells[3], 'Đơn Giá', bold=True, align='center')
        format_cell(hdr_cells[4], 'Thành Tiền', bold=True, align='center')
        
        # Điền dữ liệu bảng 2
        for i, (_, row) in enumerate(group.iterrows()):
            row_cells = bang.rows[i + 1].cells
            format_cell(row_cells[0], str(i + 1), align='center')
            format_cell(row_cells[1], str(row['tenhang']) if pd.notna(row['tenhang']) else "", align='center')
            format_cell(row_cells[2], format_number(row['soluong']), align='center')
            format_cell(row_cells[3], format_number(row['dongia']), align='right') 
            format_cell(row_cells[4], format_number(row['thanhtien']), align='right') 
            
        # Dòng: Thuế Suất GTGT 
        thue_row = bang.rows[so_mat_hang + 1].cells
        thue_row[0].merge(thue_row[3]) 
        thue_txt = f"Thuế Suất GTGT: {first_row['thuesuat'] if pd.notna(first_row['thuesuat']) else ''}"
        format_cell(thue_row[0], thue_txt, bold=True, align='center')
        # Gắn TỔNG thuế của toàn bộ các mặt hàng
        format_cell(thue_row[4], format_number(tong_thue_hop_dong), bold=True, align='right') 
        
        # Dòng: Tổng tiền thanh toán 
        tong_row = bang.rows[so_mat_hang + 2].cells
        tong_row[0].merge(tong_row[3])
        format_cell(tong_row[0], "Tổng tiền thanh toán", bold=True, align='center')
        # Gắn TỔNG tiền thanh toán của toàn bộ các mặt hàng
        format_cell(tong_row[4], format_number(tong_tien_hop_dong), bold=True, align='right') 
        # =================================================================
            
        context = {
            'sohd': so_hd_str,
            'ngay': ngay,
            'thang': thang,
            'nam': nam,
            'giaohang': giaohang,
            'thanggiaohang': thanggiaohang,
            'namgiaohang': namgiaohang,
            'benmua': str(first_row['benmua']) if pd.notna(first_row['benmua']) else "",
            'diachi': str(first_row['diachi']) if pd.notna(first_row['diachi']) else "",
            'mst': xu_ly_mst(first_row['mst']),
            'gioitinh': str(first_row['gioitinh']) if pd.notna(first_row['gioitinh']) else "",
            'daidien': str(first_row['daidien']) if pd.notna(first_row['daidien']) else "",
            
            # Chèn 2 khối subdoc bảng vào 2 vị trí
            'bang_dieu_1': bang_1_subdoc,
            'bang_hang_hoa': bang_subdoc
        }

        file_name = f"{so_hd_str}. Hợp đồng gạo - {ngay}{thang}.docx"
        output_path = os.path.join(output_dir, file_name)

        try:
            doc.render(context)
            doc.save(output_path)
            # Lưu lại đường dẫn để lát nữa "khâu" chúng nó lại
            danh_sach_file_da_tao.append(output_path) 
            print(f"✅ Đã xuất xưởng: {file_name} (Gồm {so_mat_hang} mặt hàng)")
        except Exception as e:
            print(f"❌ Lỗi ngoài ý muốn: {e}")
            return

    # =================================================================
    # TUYỆT KỸ CUỐI CÙNG: KHÂU TẤT CẢ FILE LẺ THÀNH MASTER FILE
    # =================================================================
    if danh_sach_file_da_tao:
        print("\n🔄 Đang tiến hành thuật toán gộp tất cả thành 1 file Tổng (Master File)...")
        
        # Lấy file đầu tiên làm "xương sống"
        master_doc = Document(danh_sach_file_da_tao[0])
        composer = Composer(master_doc)
        
        # Đắp lần lượt các file còn lại vào xương sống
        for file_path in danh_sach_file_da_tao[1:]:
            # Ngắt trang để mỗi hợp đồng nằm trọn vẹn ở một trang mới
            master_doc.add_page_break()
            doc_to_append = Document(file_path)
            composer.append(doc_to_append)
            
        # Đặt tên và lưu Master File
        master_file_name = f"HOP_DONG_GAO_TU_{sohd_dau}_DEN_{sohd_cuoi}.docx"
        master_file_path = os.path.join(output_dir, master_file_name)
        composer.save(master_file_path)
        
        print(f"🌟 BÙM! Đã đúc thành công bảo kiếm: {master_file_name}")

    print(f"\n🎉 XONG! Cậu mở '{output_dir}' ra xem thử Master File chuẩn bị đi in nhé!")

if __name__ == "__main__":
    tao_hop_dong_gom_nhom()
