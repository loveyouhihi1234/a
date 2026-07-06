import os
import sys
import subprocess
import warnings
from datetime import datetime, timedelta

# Tắt còi báo động (cằn nhằn) của openpyxl cho giao diện chạy sạch đẹp
warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')

# =====================================================================
# KHỐI LỆNH TỰ ĐỘNG TRANG BỊ VŨ KHÍ (PIP INSTALL)
# =====================================================================
def auto_install_vukhi(package_name, install_name=None):
    if install_name is None:
        install_name = package_name
    try:
        __import__(package_name)
    except ImportError:
        print(f"🛠 Tớ phát hiện máy cậu đang thiếu '{package_name}'...")
        print(f"⚙️ Đang tự động cài đặt luôn cho cậu. Đợi tớ 10 giây nhé!")
        subprocess.check_call([sys.executable, "-m", "pip", "install", install_name])
        print(f"✅ Trang bị vũ khí {package_name} thành công!\n")

# Tự động tải các "đồ nghề" thiết yếu
auto_install_vukhi('docxtpl')
auto_install_vukhi('pandas')
auto_install_vukhi('docx', 'python-docx')
auto_install_vukhi('openpyxl') 
auto_install_vukhi('jinja2')
# Vũ khí mới: Chuyên gia gộp file Word siêu cấp
auto_install_vukhi('docxcompose') 

import pandas as pd
import jinja2
from docxtpl import DocxTemplate
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =====================================================================
# CẤU HÌNH ĐƯỜNG DẪN FILE (Đã đổi tên theo đúng ý cậu)
# =====================================================================
DATA_FILE = '2_thong_tin_hop_dong_xe_tai.xlsx'  
WORD_TEMPLATE = '2_mau_hop_dong_xe_tai.docx' 
OUTPUT_DIR = 'Hop_Dong_Xe_Tai_Da_Xuat'

# =====================================================================
# CÁC HÀM XỬ LÝ DỮ LIỆU THÔNG MINH
# =====================================================================

def xu_ly_cmt_cccd(val):
    if pd.isna(val) or str(val).strip() == "" or str(val).lower() == 'nan':
        return ""
    s = str(val).split('.')[0].strip()
    if len(s) == 8 or len(s) == 11:
        s = "0" + s
    return s

def xu_ly_ngay_excel(val):
    if pd.isna(val) or str(val).strip() == "" or str(val).lower() == 'nan':
        return ""
    if isinstance(val, (pd.Timestamp, datetime)):
        return val.strftime("%d/%m/%Y")
    if isinstance(val, str):
        try:
            return pd.to_datetime(val).strftime("%d/%m/%Y")
        except:
            return str(val).strip()
    try:
        delta = timedelta(days=float(val))
        date_obj = datetime(1899, 12, 30) + delta
        return date_obj.strftime("%d/%m/%Y")
    except Exception:
        return str(val)

def format_number(val):
    try:
        if pd.isna(val) or val == "" or str(val).lower() == 'nan':
            return ""
        return str(int(float(val)))
    except ValueError:
        return str(val)

def format_money(val):
    try:
        if pd.isna(val) or val == "" or str(val).lower() == 'nan':
            return ""
        return f"{float(val):,.0f}".replace(',', '.')
    except ValueError:
        return str(val)

def lay_gia_tri(row, keys):
    for k in keys:
        for col in row.index:
            if k == str(col).strip():
                return row[col]
    for k in keys:
        for col in row.index:
            if k in str(col):
                return row[col]
    return ""

# =====================================================================
# BỘ CÔNG CỤ VẼ BẢNG
# =====================================================================

def format_cell(cell, text, bold=False, align='center'):
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER 
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold

def ke_vien_full(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    else:
        tblBorders.clear() 
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single') 
        border.set(qn('w:sz'), '4')       
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto') 
        tblBorders.append(border)

def autofit_to_window(table):
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000') 
    tblW.set(qn('w:type'), 'pct')

# =====================================================================
# DÂY CHUYỀN SẢN XUẤT CHÍNH
# =====================================================================
def san_xuat_hop_dong():
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    print("🚀 Khởi động dây chuyền in hợp đồng xe tải tự động (Có tính năng gộp file Master)...")

    try:
        df_raw = pd.read_excel(DATA_FILE, engine='openpyxl', header=None)
        
        header_idx = -1
        max_score = 0
        keywords = ['so', 'stt', 'số', 'bks', 'biển', 'daidien', 'đại diện', 'xe', 'oto', 'ô tô', 'trongtai', 'tải', 'cmt', 'cccd', 'ngay', 'tien', 'tiền', 'thang', 'tháng', 'tong', 'tổng']
        
        for i, row in df_raw.iterrows():
            row_text = " ".join([str(val).lower() for val in row.values if pd.notna(val)])
            score = sum(1 for k in keywords if k in row_text)
            if score > max_score:
                max_score = score
                header_idx = i
                
        if max_score < 3:
            print(f"\n❌ LỖI: Không tìm thấy dòng tiêu đề trong file {DATA_FILE}!")
            return

        print(f"🎯 Đã khóa mục tiêu dòng tiêu đề ở dòng số {header_idx + 1}. Bắt đầu sản xuất...\n")
        
        df = df_raw.iloc[header_idx + 1:].reset_index(drop=True)
        df.columns = df_raw.iloc[header_idx].astype(str).str.strip().str.lower()
        
    except FileNotFoundError:
        print(f"\n❌ Úi, tớ không tìm thấy file '{DATA_FILE}' đâu cả!")
        return
    except Exception as e:
        print(f"\n❌ Lỗi đọc file Data rồi cậu ơi: {e}")
        return

    # Mảng này để tớ lưu lại đường dẫn của tất cả các file đã xuất ra
    danh_sach_file_da_xuat = []
    thanh_cong = 0

    for index, row in df.iterrows():
        try:
            bks_raw = str(lay_gia_tri(row, ['bks', 'biển'])).strip()
            daidien_raw = str(lay_gia_tri(row, ['daidien', 'đại diện', 'tên']))
            
            if not bks_raw or bks_raw.lower() == 'nan':
                continue

            doc = DocxTemplate(WORD_TEMPLATE)

            so_raw = lay_gia_tri(row, ['so', 'stt', 'số'])
            try:
                so_str = str(int(float(so_raw))).zfill(2)
            except:
                so_str = "00"

            cmt = xu_ly_cmt_cccd(lay_gia_tri(row, ['cmt', 'cccd', 'cmnd']))
            ngaycap = xu_ly_ngay_excel(lay_gia_tri(row, ['ngaycap', 'ngày cấp', 'capngay']))
            trongtai = format_number(lay_gia_tri(row, ['trongtai', 'trọng tải', 'tải']))
            
            tien_thang = format_money(lay_gia_tri(row, ['tien', 'tiền', 'giá']))
            so_thang = format_number(lay_gia_tri(row, ['thang', 'tháng']))
            tong_tien = format_money(lay_gia_tri(row, ['tong', 'tổng', 'thành']))

            bang_subdoc = doc.new_subdoc()
            bang = bang_subdoc.add_table(rows=3, cols=5)
            ke_vien_full(bang) 
            autofit_to_window(bang)
            
            hdr = bang.rows[0].cells
            format_cell(hdr[0], 'Từ tháng', bold=True, align='center')
            format_cell(hdr[1], 'Đến tháng', bold=True, align='center')
            format_cell(hdr[2], 'Số tháng', bold=True, align='center')
            format_cell(hdr[3], 'Giá trị thuê / tháng', bold=True, align='center')
            format_cell(hdr[4], 'Thành tiền', bold=True, align='center')
            
            r1 = bang.rows[1].cells
            format_cell(r1[0], '01', align='center') 
            format_cell(r1[1], '12', align='center') 
            format_cell(r1[2], so_thang, align='center')
            format_cell(r1[3], tien_thang, align='right') 
            format_cell(r1[4], tong_tien, align='right')   
            
            r2 = bang.rows[2].cells
            r2[0].merge(r2[3]) 
            format_cell(r2[0], 'Tổng cộng:', bold=True, align='center')
            format_cell(r2[4], tong_tien, bold=True, align='right')

            context = {
                'so': so_str,
                'daidien': daidien_raw if daidien_raw.lower() != 'nan' else "",
                'diachi': str(lay_gia_tri(row, ['diachi', 'địa chỉ'])) if str(lay_gia_tri(row, ['diachi', 'địa chỉ'])).lower() != 'nan' else "",
                'cmt': cmt,
                'ngaycap': ngaycap,
                'capngay': ngaycap, 
                'congan': str(lay_gia_tri(row, ['congan', 'công an', 'nơi'])) if str(lay_gia_tri(row, ['congan', 'công an', 'nơi'])).lower() != 'nan' else "",
                'bks': bks_raw,
                'oto': str(lay_gia_tri(row, ['oto', 'ô tô', 'loại'])) if str(lay_gia_tri(row, ['oto', 'ô tô', 'loại'])).lower() != 'nan' else "",
                'trongtai': trongtai,
                'bang_thanh_toan': bang_subdoc
            }

            file_name = f"{so_str}_{bks_raw}.docx"
            output_path = os.path.join(OUTPUT_DIR, file_name)
            
            doc.render(context)
            doc.save(output_path)
            
            # Ghi chép lại file vừa tạo để tí nữa mang đi gộp
            danh_sach_file_da_xuat.append(output_path)
            print(f"✅ Đã xuất: {file_name} (Chủ xe: {context['daidien']})")
            thanh_cong += 1
            
        except jinja2.exceptions.TemplateSyntaxError as e:
            print(f"\n❌ BÁO ĐỘNG ĐỎ: Lỗi gõ thẻ bên trong file Word '{WORD_TEMPLATE}'!")
            break 
        except Exception as e:
            print(f"⚠️ Bỏ qua 1 dòng do lỗi: {e}")

    # =====================================================================
    # KHÚC CUA KHÉT LẸT: GỘP TẤT CẢ FILE LÀM MỘT CHO CẬU DỄ IN ẤN
    # =====================================================================
    if len(danh_sach_file_da_xuat) > 1:
        print("\n🔄 Bắt đầu tung ma thuật gộp file...")
        try:
            # Lấy file đầu tiên làm "trụ cột"
            master_doc = Document(danh_sach_file_da_xuat[0])
            composer = Composer(master_doc)
            
            # Quét qua các file còn lại và nối vào file "trụ cột"
            for file_path in danh_sach_file_da_xuat[1:]:
                # Tự động nhét 1 dấu ngắt trang để hợp đồng sau sang trang mới
                master_doc.add_page_break() 
                
                doc_to_append = Document(file_path)
                composer.append(doc_to_append)

            # Lưu lại siêu tác phẩm
            master_file_name = "00_Ban_In_Tat_Ca_Hop_Dong.docx"
            master_output_path = os.path.join(OUTPUT_DIR, master_file_name)
            composer.save(master_output_path)
            
            print(f"🌟 ĐÃ LUYỆN THÀNH CÔNG SIÊU TÁC PHẨM: {master_file_name} (Chỉ cần 1 click để in tất cả!)")
        except Exception as e:
            print(f"⚠️ Hơi tiếc một chút, lúc gộp file gặp lỗi nhỏ: {e}")

    if thanh_cong > 0:
        print(f"\n🎉 XONG! Tớ đã nhả ra {thanh_cong} bản hợp đồng lẻ và 1 bản Master. Cậu mở '{OUTPUT_DIR}' ra mà lấy hàng nhé!")

if __name__ == "__main__":
    san_xuat_hop_dong()