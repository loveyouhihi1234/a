# -*- coding: utf-8 -*-
import pandas as pd
from docxtpl import DocxTemplate
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from docx.oxml.ns import qn
import os
from docx import Document
from docxcompose.composer import Composer

def format_currency(value):
    """Định dạng số tiền: 33600000 -> 33.600.000"""
    try:
        if pd.isna(value) or value == '':
            return "0"
        return f"{int(value):,.0f}".replace(',', '.')
    except:
        return str(value)

def style_cell(cell, text, align_horz, is_bold=False):
    """Hàm phụ trợ: Định dạng chữ Times New Roman, cỡ 12, căn lề và dãn dòng 1.3"""
    cell.text = str(text)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
    for p in cell.paragraphs:
        p.alignment = align_horz
        p.paragraph_format.line_spacing = 1.3 
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman') 
            r.font.size = Pt(12)
            r.font.bold = is_bold

def merge_docs(file_list, output_path):
    """Hàm phụ trợ: Gộp danh sách các file Word lại thành 1 file"""
    if not file_list:
        return
    master_doc = Document(file_list[0])
    composer = Composer(master_doc)
    for file_path in file_list[1:]:
        doc_to_append = Document(file_path)
        master_doc.add_page_break() # Ngắt trang trước khi nối file mới
        composer.append(doc_to_append)
    composer.save(output_path)

def generate_documents(excel_path, hdnt_template_path, bbgn_template_path, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        
    hdnt_folder = os.path.join(output_folder, 'HopDongNguyenTac')
    bbgn_folder = os.path.join(output_folder, 'BienBanGiaoNhan')
    os.makedirs(hdnt_folder, exist_ok=True)
    os.makedirs(bbgn_folder, exist_ok=True)

    print("Đang đọc dữ liệu từ Excel...")
    df = pd.read_excel(excel_path)
    
    # Xử lý tránh đuôi .0
    df['mst'] = df['mst'].astype(str).str.replace('.0', '', regex=False)
    df['sohd'] = df['sohd'].astype(str).str.replace('.0', '', regex=False)

    # Khởi tạo 2 giỏ chứa danh sách file đã tạo
    danh_sach_hdnt = []
    danh_sach_bbgn = []

    # ==========================================
    # PHẦN 1: TẠO HỢP ĐỒNG NGUYÊN TẮC
    # ==========================================
    print("Đang xử lý Hợp Đồng Nguyên Tắc...")
    unique_invoices_per_buyer = df.groupby('benmua')['sohd'].nunique()
    buyers_need_hdnt = unique_invoices_per_buyer[unique_invoices_per_buyer >= 2].index.tolist()

    hdnt_counter = 1
    for buyer in buyers_need_hdnt:
        buyer_data = df[df['benmua'] == buyer].iloc[0]
        
        doc = DocxTemplate(hdnt_template_path)
        context_hdnt = {
            'so_hdnt': f"{hdnt_counter:02d}",
            'ky_hieu': str(buyer_data['kyhieu']) if not pd.isna(buyer_data['kyhieu']) else "",
            'ngay': f"{int(buyer_data['ngay']):02d}",
            'thang': f"{int(buyer_data['thang']):02d}",
            'nam': str(buyer_data['nam']),
            'ben_mua': str(buyer_data['benmua']),
            'dia_chi': str(buyer_data['diachi']),
            'mst': str(buyer_data['mst']),
            'gioi_tinh': str(buyer_data['gioitinh']),
            'dai_dien': str(buyer_data['daidien'])
        }
        
        doc.render(context_hdnt)
        
        ngaythang = f"{int(buyer_data['ngay']):02d}{int(buyer_data['thang']):02d}"
        safe_name = str(buyer_data['benmua']).replace('/', '_').replace('\\', '_').replace(':', '_')
        out_name = f"HDNT_{hdnt_counter:02d}_{safe_name}_{ngaythang}.docx"
        
        full_hdnt_path = os.path.join(hdnt_folder, out_name)
        doc.save(full_hdnt_path)
        danh_sach_hdnt.append(full_hdnt_path) # Ném vào giỏ HĐNT
        hdnt_counter += 1

    # ==========================================
    # PHẦN 2: TỰ VẼ BẢNG CHO BIÊN BẢN GIAO NHẬN
    # ==========================================
    print("Đang xử lý Biên Bản Giao Nhận...")
    grouped_by_invoice = df.groupby('sohd')
    
    for sohd, group in grouped_by_invoice:
        info = group.iloc[0]
        
        items = []
        for idx, row in group.reset_index().iterrows():
            items.append({
                'stt': idx + 1,
                'ten_hang': str(row['tenhang']),
                'so_luong': format_currency(row['soluong']),
                'don_gia': format_currency(row['dongia']),
                'thanh_tien': format_currency(row['thanhtien'])
            })
            
        if len(group) == 1:
            tong_thue = info['thue']
            tong_tien = info['tongtien']
        else:
            tong_thue = group['thue'].sum()
            tong_tien = group['tongtien'].sum()
            
        thue_suat_str = "" if pd.isna(info['thuesuat']) else str(info['thuesuat'])
        
        doc_bbgn = DocxTemplate(bbgn_template_path)
        context_bbgn = {
            'ngay': f"{int(info['ngay']):02d}",
            'thang': f"{int(info['thang']):02d}",
            'nam': str(info['nam']),
            'ben_mua': str(info['benmua']),
            'dia_chi': str(info['diachi']),
            'mst': str(info['mst']),
            'gioi_tinh': str(info['gioitinh']),
            'dai_dien': str(info['daidien'])
        }
        
        doc_bbgn.render(context_bbgn)
        
        for p in doc_bbgn.docx.paragraphs:
            if '[CHEN_BANG_VAO_DAY]' in p.text:
                table = doc_bbgn.docx.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                hdr_cells = table.rows[0].cells
                headers = ['STT', 'Tên hàng hóa', 'Số lượng (Kg)\n(±10%)', 'Đơn Giá\n(Đồng/Kg)', 'Thành Tiền\n(VNĐ)']
                for i, header_text in enumerate(headers):
                    style_cell(hdr_cells[i], header_text, WD_ALIGN_PARAGRAPH.CENTER, is_bold=True)
                
                for item in items:
                    row_cells = table.add_row().cells
                    style_cell(row_cells[0], item['stt'], WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell(row_cells[1], item['ten_hang'], WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell(row_cells[2], item['so_luong'], WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell(row_cells[3], item['don_gia'], WD_ALIGN_PARAGRAPH.RIGHT)
                    style_cell(row_cells[4], item['thanh_tien'], WD_ALIGN_PARAGRAPH.RIGHT)
                
                tax_row = table.add_row().cells
                tax_row[0].merge(tax_row[3])
                style_cell(tax_row[0], f"Thuế Suất GTGT: {thue_suat_str}", WD_ALIGN_PARAGRAPH.CENTER)
                style_cell(tax_row[4], format_currency(tong_thue), WD_ALIGN_PARAGRAPH.RIGHT)
                
                total_row = table.add_row().cells
                total_row[0].merge(total_row[3])
                style_cell(total_row[0], "Tổng tiền thanh toán", WD_ALIGN_PARAGRAPH.CENTER)
                style_cell(total_row[4], format_currency(tong_tien), WD_ALIGN_PARAGRAPH.RIGHT)
                
                p._p.addprevious(table._tbl)
                p._element.getparent().remove(p._element)
                break

        ngaythang = f"{int(info['ngay']):02d}{int(info['thang']):02d}"
        safe_name = str(info['benmua']).replace('/', '_').replace('\\', '_').replace(':', '_')
        out_bbgn_name = f"BBGN_HD{sohd}_{safe_name}_{ngaythang}.docx"
        
        full_bbgn_path = os.path.join(bbgn_folder, out_bbgn_name)
        doc_bbgn.save(full_bbgn_path)
        danh_sach_bbgn.append(full_bbgn_path) # Ném vào giỏ BBGN

    # ==========================================
    # PHẦN 3: ĐÓNG QUYỂN (GỘP FILE)
    # ==========================================
    print("Đang tiến hành đóng quyển các văn bản...")
    
    # 1. Gộp tất cả Hợp đồng nguyên tắc
    if danh_sach_hdnt:
        file_tong_hdnt = os.path.join(output_folder, '1_Tong_Hop_Hop_Dong_Nguyen_Tac.docx')
        merge_docs(danh_sach_hdnt, file_tong_hdnt)
        print(f"-> Đã gộp xong: {file_tong_hdnt}")

    # 2. Gộp tất cả Biên bản giao nhận
    if danh_sach_bbgn:
        file_tong_bbgn = os.path.join(output_folder, '2_Tong_Hop_Bien_Ban_Giao_Nhan.docx')
        merge_docs(danh_sach_bbgn, file_tong_bbgn)
        print(f"-> Đã gộp xong: {file_tong_bbgn}")

    # 3. Gộp CẢ Hợp đồng và Biên bản vào làm một
    danh_sach_tong_hop = danh_sach_hdnt + danh_sach_bbgn
    if danh_sach_tong_hop:
        file_tong_ca_hai = os.path.join(output_folder, '3_Tong_Hop_Tat_Ca_HDNT_va_BBGN.docx')
        merge_docs(danh_sach_tong_hop, file_tong_ca_hai)
        print(f"-> Đã gộp xong: {file_tong_ca_hai}")

    print("🎉 Quá mượt! Toàn bộ 3 file tổng hợp đã nằm gọn gàng trong thư mục đầu ra.")

if __name__ == '__main__':
    EXCEL_FILE = '3_danh_sach_hoa_don_gao.xlsx'          
    TEMPLATE_HDNT = '3_mau_hopdongnguyentac.docx'    
    TEMPLATE_BBGN = '3_mau_bienbangiaonhan.docx'     
    OUTPUT_DIR = 'Xuat_File_Tu_Dong'               
    
    generate_documents(EXCEL_FILE, TEMPLATE_HDNT, TEMPLATE_BBGN, OUTPUT_DIR)