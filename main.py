import os
import io
import zipfile
import tempfile
import pandas as pd
import jinja2
from datetime import datetime, timedelta
from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

from docxtpl import DocxTemplate
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docxcompose.composer import Composer

app = FastAPI(title="Hệ thống Tự động hóa Hợp đồng")

# Mở khóa bảo mật CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =====================================================================
# CÁC HÀM XỬ LÝ DỮ LIỆU CHUNG
# =====================================================================
def xu_ly_cmt_cccd(val):
    if pd.isna(val) or str(val).strip() == "" or str(val).lower() == 'nan': return ""
    s = str(val).split('.')[0].strip()
    if len(s) == 8 or len(s) == 11: s = "0" + s
    return s

def xu_ly_ngay_excel(val):
    if pd.isna(val) or str(val).strip() == "" or str(val).lower() == 'nan': return ""
    if isinstance(val, (pd.Timestamp, datetime)): return val.strftime("%d/%m/%Y")
    if isinstance(val, str):
        try: return pd.to_datetime(val).strftime("%d/%m/%Y")
        except: return str(val).strip()
    try:
        delta = timedelta(days=float(val))
        date_obj = datetime(1899, 12, 30) + delta
        return date_obj.strftime("%d/%m/%Y")
    except Exception: return str(val)

def format_number(val):
    try:
        if pd.isna(val) or val == "" or str(val).lower() == 'nan': return ""
        return f"{float(val):,.0f}".replace(',', '.')
    except ValueError: return str(val)

def xu_ly_mst(val):
    if pd.isna(val) or str(val).strip() == "": return ""
    s = str(val).strip()
    if s.endswith('.0'): s = s[:-2]
    if s.isdigit():
        if len(s) <= 10: return s.zfill(10)
        elif len(s) <= 13:
            s_13 = s.zfill(13)
            return f"{s_13[:10]}-{s_13[10:]}"
    return s

def lay_gia_tri(row, keys):
    for k in keys:
        for col in row.index:
            if k == str(col).strip(): return row[col]
    for k in keys:
        for col in row.index:
            if k in str(col): return row[col]
    return ""

def format_cell(cell, text, bold=False, align='center'):
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'left': p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold

def autofit_to_window(table):
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')

def tao_file_zip(thu_muc_nguon, ten_file_zip_ao):
    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(thu_muc_nguon):
            for file in files:
                file_path = os.path.join(root, file)
                zf.write(file_path, os.path.relpath(file_path, thu_muc_nguon))
    memory_file.seek(0)
    return StreamingResponse(
        memory_file, 
        media_type="application/zip", 
        headers={"Content-Disposition": f"attachment; filename={ten_file_zip_ao}"}
    )

@app.get("/")
def read_root():
    return {"message": "Hệ thống Tự động hóa Hợp đồng đang bay mượt trên đám mây!"}

# =====================================================================
# API 1: HỢP ĐỒNG XE TẢI
# =====================================================================
def ke_vien_full(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    else: tblBorders.clear()
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)

@app.post("/api/hop-dong-xe-tai")
async def tao_hop_dong_xe_tai(excel_file: UploadFile = File(...), word_template: UploadFile = File(...)):
    excel_data = await excel_file.read()
    word_data = await word_template.read()
    
    df_raw = pd.read_excel(io.BytesIO(excel_data), engine='openpyxl', header=None)
    
    header_idx = -1
    max_score = 0
    keywords = ['so', 'stt', 'số', 'bks', 'biển', 'daidien', 'đại diện', 'xe', 'oto', 'ô tô', 'trongtai', 'tải', 'cmt', 'cccd', 'ngay', 'tien', 'tiền', 'thang', 'tháng', 'tong', 'tổng']
    
    for i, row in df_raw.iterrows():
        row_text = " ".join([str(val).lower() for val in row.values if pd.notna(val)])
        score = sum(1 for k in keywords if k in row_text)
        if score > max_score:
            max_score = score
            header_idx = i

    if max_score < 3: 
        return {"error": "Không tìm thấy dòng tiêu đề hợp lệ."}
    
    df = df_raw.iloc[header_idx + 1:].reset_index(drop=True)
    df.columns = df_raw.iloc[header_idx].astype(str).str.strip().str.lower()

    with tempfile.TemporaryDirectory() as temp_dir:
        for index, row in df.iterrows():
            bks_raw = str(lay_gia_tri(row, ['bks', 'biển'])).strip()
            daidien_raw = str(lay_gia_tri(row, ['daidien', 'đại diện', 'tên']))
            if not bks_raw or bks_raw.lower() == 'nan': continue

            doc = DocxTemplate(io.BytesIO(word_data))
            
            so_raw = lay_gia_tri(row, ['so', 'stt', 'số'])
            try: so_str = str(int(float(so_raw))).zfill(2)
            except: so_str = "00"

            cmt = xu_ly_cmt_cccd(lay_gia_tri(row, ['cmt', 'cccd', 'cmnd']))
            ngaycap = xu_ly_ngay_excel(lay_gia_tri(row, ['ngaycap', 'ngày cấp', 'capngay']))
            trongtai = format_number(lay_gia_tri(row, ['trongtai', 'trọng tải', 'tải']))
            tien_thang = format_number(lay_gia_tri(row, ['tien', 'tiền', 'giá']))
            so_thang = format_number(lay_gia_tri(row, ['thang', 'tháng']))
            tong_tien = format_number(lay_gia_tri(row, ['tong', 'tổng', 'thành']))

            bang_subdoc = doc.new_subdoc()
            bang = bang_subdoc.add_table(rows=3, cols=5)
            ke_vien_full(bang)
            autofit_to_window(bang)

            hdr = bang.rows[0].cells
            format_cell(hdr[0], 'Từ tháng', bold=True)
            format_cell(hdr[1], 'Đến tháng', bold=True)
            format_cell(hdr[2], 'Số tháng', bold=True)
            format_cell(hdr[3], 'Giá trị thuê / tháng', bold=True)
            format_cell(hdr[4], 'Thành tiền', bold=True)

            r1 = bang.rows[1].cells
            format_cell(r1[0], '01')
            format_cell(r1[1], '12')
            format_cell(r1[2], so_thang)
            format_cell(r1[3], tien_thang, align='right')
            format_cell(r1[4], tong_tien, align='right')

            r2 = bang.rows[2].cells
            r2[0].merge(r2[3])
            format_cell(r2[0], 'Tổng cộng:', bold=True)
            format_cell(r2[4], tong_tien, bold=True, align='right')

            context = {
                'so': so_str,
                'daidien': daidien_raw if daidien_raw.lower() != 'nan' else "",
                'diachi': str(lay_gia_tri(row, ['diachi', 'địa chỉ'])) if str(lay_gia_tri(row, ['diachi', 'địa chỉ'])).lower() != 'nan' else "",
                'cmt': cmt,
                'ngaycap': ngaycap,
                'capngay': ngaycap,
                'congan': str(lay_gia_tri(row, ['congan', 'công an', 'nơi'])) if str(lay_gia_tri(row, ['congan', 'công an', 'nơi'])).lower() != 'nan' else "",
                'bks': bks_raw,
                'oto': str(lay_gia_tri(row, ['oto', 'ô tô', 'loại'])) if str(lay_gia_tri(row, ['oto', 'ô tô', 'loại'])).lower() != 'nan' else "",
                'trongtai': trongtai,
                'bang_thanh_toan': bang_subdoc
            }

            file_name = f"{so_str}_{bks_raw}.docx"
            doc.render(context)
            doc.save(os.path.join(temp_dir, file_name))

        return tao_file_zip(temp_dir, "HopDongXeTai_ThanhPham.zip")

# =====================================================================
# API 2: HỢP ĐỒNG GẠO ĐƠN THUẦN
# =====================================================================
def ve_vien_ngang(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None: tblBorders.clear()
    else:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    for border_name in ['insideH']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') 
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)

@app.post("/api/hop-dong-gao")
async def tao_hop_dong_gao(excel_file: UploadFile = File(...), word_template: UploadFile = File(...)):
    excel_data = await excel_file.read()
    word_data = await word_template.read()
    
    df = pd.read_excel(io.BytesIO(excel_data))
    df = df.dropna(subset=['sohd']) 
    nhom_hop_dong = df.groupby('sohd')

    with tempfile.TemporaryDirectory() as temp_dir:
        for sohd, group in nhom_hop_dong:
            doc = DocxTemplate(io.BytesIO(word_data))
            
            first_row = group.iloc[0]
            so_hd_str = str(sohd).replace('.0', '').zfill(2)
            ngay = str(first_row['ngay']).replace('.0', '').zfill(2)
            thang = str(first_row['thang']).replace('.0', '').zfill(2)
            nam = str(first_row['nam']).replace('.0', '')
            
            try:
                ngay_ky_obj = datetime(int(nam), int(thang), int(ngay))
                ngay_giao_obj = ngay_ky_obj + timedelta(days=2)
                giaohang = ngay_giao_obj.strftime("%d")
                thanggiaohang = ngay_giao_obj.strftime("%m")
                namgiaohang = ngay_giao_obj.strftime("%Y")
            except Exception:
                giaohang, thanggiaohang, namgiaohang = "...", "...", "..."
                
            so_mat_hang = len(group)

            bang_1_subdoc = doc.new_subdoc()
            bang_1 = bang_1_subdoc.add_table(rows=so_mat_hang + 1, cols=4)
            ve_vien_ngang(bang_1)
            autofit_to_window(bang_1)
            
            hdr1 = bang_1.rows[0].cells
            format_cell(hdr1[0], 'Tên hàng', bold=True)
            format_cell(hdr1[1], 'Số lượng', bold=True)
            format_cell(hdr1[2], 'Đơn giá', bold=True)
            format_cell(hdr1[3], 'Mặt hàng', bold=True)
            
            for i, (_, row) in enumerate(group.iterrows()):
                r1 = bang_1.rows[i + 1].cells
                format_cell(r1[0], str(row['tenhang']) if pd.notna(row['tenhang']) else "")
                format_cell(r1[1], f"{format_number(row['soluong'])} kg")
                format_cell(r1[2], f"{format_number(row['dongia'])} đ/kg", align='right') 
                format_cell(r1[3], str(row['mathang']) if pd.notna(row['mathang']) else "")

            bang_subdoc = doc.new_subdoc()
            bang = bang_subdoc.add_table(rows=so_mat_hang + 3, cols=5)
            bang.style = 'Table Grid'
            autofit_to_window(bang) 
            
            hdr_cells = bang.rows[0].cells
            format_cell(hdr_cells[0], 'STT', bold=True)
            format_cell(hdr_cells[1], 'Tên hàng hóa', bold=True)
            format_cell(hdr_cells[2], 'Số lượng (Kg)', bold=True)
            format_cell(hdr_cells[3], 'Đơn Giá', bold=True)
            format_cell(hdr_cells[4], 'Thành Tiền', bold=True)
            
            for i, (_, row) in enumerate(group.iterrows()):
                row_cells = bang.rows[i + 1].cells
                format_cell(row_cells[0], str(i + 1))
                format_cell(row_cells[1], str(row['tenhang']) if pd.notna(row['tenhang']) else "")
                format_cell(row_cells[2], format_number(row['soluong']))
                format_cell(row_cells[3], format_number(row['dongia']), align='right') 
                format_cell(row_cells[4], format_number(row['thanhtien']), align='right') 
                
            thue_row = bang.rows[so_mat_hang + 1].cells
            thue_row[0].merge(thue_row[3]) 
            thue_txt = f"Thuế Suất GTGT: {first_row['thuesuat'] if pd.notna(first_row['thuesuat']) else ''}"
            format_cell(thue_row[0], thue_txt, bold=True)
            format_cell(thue_row[4], format_number(first_row['thue']), bold=True, align='right') 
            
            tong_row = bang.rows[so_mat_hang + 2].cells
            tong_row[0].merge(tong_row[3])
            format_cell(tong_row[0], "Tổng tiền thanh toán", bold=True)
            format_cell(tong_row[4], format_number(first_row['tongtien']), bold=True, align='right') 
                
            context = {
                'sohd': so_hd_str,
                'ngay': ngay, 'thang': thang, 'nam': nam,
                'giaohang': giaohang, 'thanggiaohang': thanggiaohang, 'namgiaohang': namgiaohang,
                'benmua': str(first_row['benmua']) if pd.notna(first_row['benmua']) else "",
                'diachi': str(first_row['diachi']) if pd.notna(first_row['diachi']) else "",
                'mst': xu_ly_mst(first_row['mst']),
                'gioitinh': str(first_row['gioitinh']) if pd.notna(first_row['gioitinh']) else "",
                'daidien': str(first_row['daidien']) if pd.notna(first_row['daidien']) else "",
                'bang_dieu_1': bang_1_subdoc,
                'bang_hang_hoa': bang_subdoc
            }

            file_name = f"{so_hd_str}. Hợp đồng gạo - {ngay}{thang}.docx"
            doc.render(context)
            doc.save(os.path.join(temp_dir, file_name))

        return tao_file_zip(temp_dir, "HopDongGao_ThanhPham.zip")


# =====================================================================
# API 3: HỢP ĐỒNG NGUYÊN TẮC GẠO (HDNT & BBGN) + ĐÓNG QUYỂN
# =====================================================================
def style_cell_api3(cell, text, align_horz, is_bold=False):
    cell.text = str(text)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
    for p in cell.paragraphs:
        p.alignment = align_horz
        p.paragraph_format.line_spacing = 1.3 
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman') 
            r.font.size = Pt(12)
            r.font.bold = is_bold

def merge_docs(file_list, output_path):
    if not file_list: return
    master_doc = Document(file_list[0])
    composer = Composer(master_doc)
    for file_path in file_list[1:]:
        doc_to_append = Document(file_path)
        master_doc.add_page_break()
        composer.append(doc_to_append)
    composer.save(output_path)

@app.post("/api/hop-dong-nguyen-tac-gao")
async def tao_hop_dong_nguyen_tac(
    excel_file: UploadFile = File(...),
    hdnt_template: UploadFile = File(...),
    bbgn_template: UploadFile = File(...)
):
    with tempfile.TemporaryDirectory() as temp_dir:
        # Ghi template ra đĩa tạm để dùng cho docxtpl
        hdnt_path = os.path.join(temp_dir, "temp_hdnt.docx")
        with open(hdnt_path, "wb") as f: f.write(await hdnt_template.read())
        
        bbgn_path = os.path.join(temp_dir, "temp_bbgn.docx")
        with open(bbgn_path, "wb") as f: f.write(await bbgn_template.read())

        # Đọc dữ liệu Excel
        excel_data = await excel_file.read()
        df = pd.read_excel(io.BytesIO(excel_data))
        df['mst'] = df['mst'].astype(str).str.replace('.0', '', regex=False)
        df['sohd'] = df['sohd'].astype(str).str.replace('.0', '', regex=False)

        hdnt_folder = os.path.join(temp_dir, 'HopDongNguyenTac')
        bbgn_folder = os.path.join(temp_dir, 'BienBanGiaoNhan')
        os.makedirs(hdnt_folder, exist_ok=True)
        os.makedirs(bbgn_folder, exist_ok=True)

        danh_sach_hdnt = []
        danh_sach_bbgn = []

        # PHẦN 1: HỢP ĐỒNG NGUYÊN TẮC
        unique_invoices_per_buyer = df.groupby('benmua')['sohd'].nunique()
        buyers_need_hdnt = unique_invoices_per_buyer[unique_invoices_per_buyer >= 2].index.tolist()

        hdnt_counter = 1
        for buyer in buyers_need_hdnt:
            buyer_data = df[df['benmua'] == buyer].iloc[0]
            doc = DocxTemplate(hdnt_path)
            
            context_hdnt = {
                'so_hdnt': f"{hdnt_counter:02d}",
                'ky_hieu': str(buyer_data['kyhieu']) if not pd.isna(buyer_data['kyhieu']) else "",
                'ngay': f"{int(buyer_data['ngay']):02d}",
                'thang': f"{int(buyer_data['thang']):02d}",
                'nam': str(buyer_data['nam']),
                'ben_mua': str(buyer_data['benmua']),
                'dia_chi': str(buyer_data['diachi']),
                'mst': str(buyer_data['mst']),
                'gioi_tinh': str(buyer_data['gioitinh']),
                'dai_dien': str(buyer_data['daidien'])
            }
            doc.render(context_hdnt)
            
            ngaythang = f"{int(buyer_data['ngay']):02d}{int(buyer_data['thang']):02d}"
            safe_name = str(buyer_data['benmua']).replace('/', '_').replace('\\', '_').replace(':', '_')
            out_name = f"HDNT_{hdnt_counter:02d}_{safe_name}_{ngaythang}.docx"
            
            full_hdnt_path = os.path.join(hdnt_folder, out_name)
            doc.save(full_hdnt_path)
            danh_sach_hdnt.append(full_hdnt_path)
            hdnt_counter += 1

        # PHẦN 2: BIÊN BẢN GIAO NHẬN
        grouped_by_invoice = df.groupby('sohd')
        for sohd, group in grouped_by_invoice:
            info = group.iloc[0]
            
            items = []
            for idx, row in group.reset_index().iterrows():
                items.append({
                    'stt': idx + 1,
                    'ten_hang': str(row['tenhang']),
                    'so_luong': format_number(row['soluong']),
                    'don_gia': format_number(row['dongia']),
                    'thanh_tien': format_number(row['thanhtien'])
                })
                
            if len(group) == 1:
                tong_thue = info['thue']
                tong_tien = info['tongtien']
            else:
                tong_thue = group['thue'].sum()
                tong_tien = group['tongtien'].sum()
                
            thue_suat_str = "" if pd.isna(info['thuesuat']) else str(info['thuesuat'])
            
            doc_bbgn = DocxTemplate(bbgn_path)
            context_bbgn = {
                'ngay': f"{int(info['ngay']):02d}",
                'thang': f"{int(info['thang']):02d}",
                'nam': str(info['nam']),
                'ben_mua': str(info['benmua']),
                'dia_chi': str(info['diachi']),
                'mst': str(info['mst']),
                'gioi_tinh': str(info['gioitinh']),
                'dai_dien': str(info['daidien'])
            }
            doc_bbgn.render(context_bbgn)
            
            for p in doc_bbgn.docx.paragraphs:
                if '[CHEN_BANG_VAO_DAY]' in p.text:
                    table = doc_bbgn.docx.add_table(rows=1, cols=5)
                    table.style = 'Table Grid'
                    
                    hdr_cells = table.rows[0].cells
                    headers = ['STT', 'Tên hàng hóa', 'Số lượng (Kg)\n(±10%)', 'Đơn Giá\n(Đồng/Kg)', 'Thành Tiền\n(VNĐ)']
                    for i, header_text in enumerate(headers):
                        style_cell_api3(hdr_cells[i], header_text, WD_ALIGN_PARAGRAPH.CENTER, is_bold=True)
                    
                    for item in items:
                        row_cells = table.add_row().cells
                        style_cell_api3(row_cells[0], item['stt'], WD_ALIGN_PARAGRAPH.CENTER)
                        style_cell_api3(row_cells[1], item['ten_hang'], WD_ALIGN_PARAGRAPH.CENTER)
                        style_cell_api3(row_cells[2], item['so_luong'], WD_ALIGN_PARAGRAPH.CENTER)
                        style_cell_api3(row_cells[3], item['don_gia'], WD_ALIGN_PARAGRAPH.RIGHT)
                        style_cell_api3(row_cells[4], item['thanh_tien'], WD_ALIGN_PARAGRAPH.RIGHT)
                    
                    tax_row = table.add_row().cells
                    tax_row[0].merge(tax_row[3])
                    style_cell_api3(tax_row[0], f"Thuế Suất GTGT: {thue_suat_str}", WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_api3(tax_row[4], format_number(tong_thue), WD_ALIGN_PARAGRAPH.RIGHT)
                    
                    total_row = table.add_row().cells
                    total_row[0].merge(total_row[3])
                    style_cell_api3(total_row[0], "Tổng tiền thanh toán", WD_ALIGN_PARAGRAPH.CENTER)
                    style_cell_api3(total_row[4], format_number(tong_tien), WD_ALIGN_PARAGRAPH.RIGHT)
                    
                    p._p.addprevious(table._tbl)
                    p._element.getparent().remove(p._element)
                    break

            ngaythang = f"{int(info['ngay']):02d}{int(info['thang']):02d}"
            safe_name = str(info['benmua']).replace('/', '_').replace('\\', '_').replace(':', '_')
            out_bbgn_name = f"BBGN_HD{sohd}_{safe_name}_{ngaythang}.docx"
            
            full_bbgn_path = os.path.join(bbgn_folder, out_bbgn_name)
            doc_bbgn.save(full_bbgn_path)
            danh_sach_bbgn.append(full_bbgn_path)

        # PHẦN 3: ĐÓNG QUYỂN
        if danh_sach_hdnt:
            merge_docs(danh_sach_hdnt, os.path.join(temp_dir, '1_Tong_Hop_Hop_Dong_Nguyen_Tac.docx'))
        if danh_sach_bbgn:
            merge_docs(danh_sach_bbgn, os.path.join(temp_dir, '2_Tong_Hop_Bien_Ban_Giao_Nhan.docx'))
        
        danh_sach_tong_hop = danh_sach_hdnt + danh_sach_bbgn
        if danh_sach_tong_hop:
            merge_docs(danh_sach_tong_hop, os.path.join(temp_dir, '3_Tong_Hop_Tat_Ca_HDNT_va_BBGN.docx'))

        # Dọn 2 file mẫu vừa lưu tạm
        os.remove(hdnt_path)
        os.remove(bbgn_path)

        return tao_file_zip(temp_dir, "KetQua_HDNT_Gao.zip")